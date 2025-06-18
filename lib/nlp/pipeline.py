"""
Enterprise-Grade NLP Pipeline Classes
Built with production-ready components and comprehensive error handling
"""

import asyncio
import logging
import re
import hashlib
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Tuple, Union, Any
from pathlib import Path
from datetime import datetime
import json

# Core NLP and ML libraries
import spacy
import torch
import numpy as np
import pandas as pd
from transformers import (
    AutoTokenizer, AutoModelForSequenceClassification,
    AutoModelForTokenClassification, pipeline
)
from sentence_transformers import SentenceTransformer
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import WordNetLemmatizer
from textstat import flesch_reading_ease, flesch_kincaid_grade, automated_readability_index

# Document processing
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pdfplumber
from bs4 import BeautifulSoup
import magic
import chardet

# Feature engineering and ML
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import StandardScaler
import gensim
from gensim import corpora
import bertopic
from bertopic import BERTopic

# Utilities
from langdetect import detect, DetectorFactory
import textacy
from textacy import preprocessing
import vaderSentiment
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import yake
from collections import Counter
import concurrent.futures
from functools import lru_cache, wraps
import time

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Ensure consistent language detection
DetectorFactory.seed = 0

@dataclass
class ProcessedDocument:
    """Data class for processed documents"""
    id: str
    original_text: str
    cleaned_text: str
    language: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    sentences: List[str] = field(default_factory=list)
    tokens: List[str] = field(default_factory=list)
    structure: Dict[str, Any] = field(default_factory=dict)
    processing_time: float = 0.0
    errors: List[str] = field(default_factory=list)

@dataclass
class NLPResults:
    """Data class for NLP processing results"""
    document_id: str
    embeddings: np.ndarray
    classification: Dict[str, Any]
    entities: List[Dict[str, Any]]
    sentiment: Dict[str, Any]
    topics: List[Dict[str, Any]]
    keywords: List[str]
    language: str
    quality_metrics: Dict[str, Any]
    processing_metadata: Dict[str, Any] = field(default_factory=dict)

def timing_decorator(func):
    """Decorator to measure function execution time"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        logger.info(f"{func.__name__} executed in {end_time - start_time:.2f} seconds")
        return result
    return wrapper

def error_handler(func):
    """Decorator for comprehensive error handling"""
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            logger.error(f"Error in {func.__name__}: {str(e)}")
            raise
    return wrapper

class DocumentPreprocessor:
    """Enterprise-grade document preprocessing with multi-format support"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.supported_formats = {'.pdf', '.docx', '.txt', '.html', '.htm', '.md'}
        
        # Initialize NLP tools
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy English model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
            
        # Download required NLTK data
        self._download_nltk_data()
        
        # Initialize text processing tools
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        
        # Language detection cache
        self._language_cache = {}
        
    def _download_nltk_data(self):
        """Download required NLTK datasets"""
        nltk_downloads = ['punkt', 'stopwords', 'wordnet', 'averaged_perceptron_tagger']
        for dataset in nltk_downloads:
            try:
                nltk.data.find(f'tokenizers/{dataset}')
            except LookupError:
                nltk.download(dataset, quiet=True)
    
    @timing_decorator
    @error_handler
    def process(self, document_input: Union[str, Path, bytes], 
                document_id: Optional[str] = None) -> ProcessedDocument:
        """
        Main processing method that handles various input types
        
        Args:
            document_input: File path, raw text, or bytes
            document_id: Optional document identifier
            
        Returns:
            ProcessedDocument with all preprocessing completed
        """
        start_time = time.time()
        
        if document_id is None:
            document_id = self._generate_document_id(document_input)
            
        # Extract text from various sources
        raw_text, metadata = self._extract_text(document_input)
        
        # Detect language
        language = self._detect_language(raw_text)
        
        # Clean and normalize text
        cleaned_text = self._clean_text(raw_text, language)
        
        # Tokenization and sentence segmentation
        sentences = self._segment_sentences(cleaned_text, language)
        tokens = self._tokenize_text(cleaned_text, language)
        
        # Structure analysis
        structure = self._analyze_structure(raw_text, cleaned_text)
        
        processing_time = time.time() - start_time
        
        return ProcessedDocument(
            id=document_id,
            original_text=raw_text,
            cleaned_text=cleaned_text,
            language=language,
            metadata=metadata,
            sentences=sentences,
            tokens=tokens,
            structure=structure,
            processing_time=processing_time
        )
    
    def _generate_document_id(self, document_input: Union[str, Path, bytes]) -> str:
        """Generate unique document ID based on content"""
        if isinstance(document_input, (str, Path)):
            content = str(document_input)
        else:
            content = str(document_input)[:1000]  # Use first 1000 chars for ID
        
        return hashlib.md5(content.encode()).hexdigest()
    
    def _extract_text(self, document_input: Union[str, Path, bytes]) -> Tuple[str, Dict[str, Any]]:
        """Extract text from various document formats"""
        metadata = {}
        
        if isinstance(document_input, str) and not Path(document_input).exists():
            # Raw text input
            return document_input, {"source": "raw_text"}
        
        file_path = Path(document_input)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect file type
        file_type = magic.from_file(str(file_path), mime=True)
        metadata["file_type"] = file_type
        metadata["file_size"] = file_path.stat().st_size
        metadata["file_path"] = str(file_path)
        
        # Extract based on file type
        if file_path.suffix.lower() == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_path.suffix.lower() == '.docx':
            text = self._extract_docx_text(file_path)
        elif file_path.suffix.lower() in {'.txt', '.md'}:
            text = self._extract_text_file(file_path)
        elif file_path.suffix.lower() in {'.html', '.htm'}:
            text = self._extract_html_text(file_path)
        else:
            # Fallback to text extraction
            text = self._extract_text_file(file_path)
        
        return text, metadata
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Try pdfplumber first (better for tables)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
            
        # Fallback to PyMuPDF
        if not text.strip():
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                raise
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text files with encoding detection"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise
    
    def _extract_html_text(self, file_path: Path) -> str:
        """Extract text from HTML files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text()
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            raise
    
    @lru_cache(maxsize=1000)
    def _detect_language(self, text: str) -> str:
        """Detect document language with caching"""
        if not text.strip():
            return "unknown"
        
        try:
            # Use first 1000 characters for detection
            sample_text = text[:1000]
            detected_lang = detect(sample_text)
            return detected_lang
        except Exception as e:
            logger.warning(f"Language detection failed: {e}")
            return "en"  # Default to English
    
    def _clean_text(self, text: str, language: str) -> str:
        """Comprehensive text cleaning and normalization"""
        if not text:
            return ""
        
        # Use textacy for advanced preprocessing
        text = preprocessing.normalize.hyphenated_words(text)
        text = preprocessing.normalize.quotation_marks(text)
        text = preprocessing.normalize.unicode(text)
        text = preprocessing.remove.accents(text)
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        # Remove special characters but keep sentence structure
        text = re.sub(r '[^\w\s\.\?\!\,\;\:\-\(\)\[\]\{\}\"\']+', ' ', text)
        
        return text.strip()
    
    def _segment_sentences(self, text: str, language: str) -> List[str]:
        """Advanced sentence segmentation"""
        if not text:
            return []
        
        if self.nlp and language == 'en':
            # Use spaCy for better segmentation
            doc = self.nlp(text)
            return [sent.text.strip() for sent in doc.sents if sent.text.strip()]
        else:
            # Fallback to NLTK
            return [sent.strip() for sent in sent_tokenize(text) if sent.strip()]
    
    def _tokenize_text(self, text: str, language: str) -> List[str]:
        """Advanced tokenization with lemmatization"""
        if not text:
            return []
        
        if self.nlp and language == 'en':
            # Use spaCy for advanced tokenization
            doc = self.nlp(text)
            tokens = []
            for token in doc:
                if not token.is_stop and not token.is_punct and not token.is_space:
                    tokens.append(token.lemma_.lower())
            return tokens
        else:
            # Fallback to NLTK
            tokens = word_tokenize(text.lower())
            tokens = [self.lemmatizer.lemmatize(token) for token in tokens 
                     if token not in self.stop_words and token.isalpha()]
            return tokens
    
    def _analyze_structure(self, raw_text: str, cleaned_text: str) -> Dict[str, Any]:
        """Analyze document structure and formatting"""
        structure = {
            "char_count": len(raw_text),
            "word_count": len(cleaned_text.split()),
            "sentence_count": len(sent_tokenize(cleaned_text)),
            "paragraph_count": len([p for p in raw_text.split('\n\n') if p.strip()]),
            "avg_sentence_length": 0,
            "avg_word_length": 0,
            "question_count": raw_text.count('?'),
            "exclamation_count": raw_text.count('!'),
            "has_headers": bool(re.search(r'^#{1,6}\s+', raw_text, re.MULTILINE)),
            "has_bullet_points": bool(re.search(r'^\s*[\*\-\+]\s+', raw_text, re.MULTILINE)),
            "has_numbered_lists": bool(re.search(r'^\s*\d+\.\s+', raw_text, re.MULTILINE)),
            "uppercase_ratio": sum(1 for c in cleaned_text if c.isupper()) / len(cleaned_text) if cleaned_text else 0
        }
        
        # Calculate averages
        words = cleaned_text.split()
        if words:
            structure["avg_word_length"] = sum(len(word) for word in words) / len(words)
        
        sentences = sent_tokenize(cleaned_text)
        if sentences:
            structure["avg_sentence_length"] = sum(len(sent.split()) for sent in sentences) / len(sentences)
        
        return structure


class NLPProcessor:
    """Enterprise NLP processing engine with multiple model support"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        logger.info(f"Using device: {self.device}")
        
        # Initialize models
        self._load_models()
        
        # Initialize processing components
        self.sentiment_analyzer = SentimentIntensityAnalyzer()
        self.keyword_extractor = yake.KeywordExtractor(
            lan="en", n=3, dedupLim=0.7, top=20
        )
        
    def _load_models(self):
        """Load and initialize all required models"""
        try:
            # Sentence transformer for embeddings
            self.sentence_transformer = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Load spaCy model
            self.nlp = spacy.load("en_core_web_sm")
            
            # Load transformers models
            self.tokenizer = AutoTokenizer.from_pretrained("distilbert-base-uncased")
            
            # Sentiment model
            self.sentiment_model = pipeline(
                "sentiment-analysis",
                model="cardiffnlp/twitter-roberta-base-sentiment-latest",
                device=0 if self.device.type == "cuda" else -1
            )
            
            # Topic modeling
            self.topic_model = BERTopic(
                embedding_model=self.sentence_transformer,
                verbose=False
            )
            
            logger.info("All models loaded successfully")
            
        except Exception as e:
            logger.error(f"Model loading failed: {e}")
            raise
    
    @timing_decorator
    @error_handler
    def process_document(self, document: ProcessedDocument) -> NLPResults:
        """
        Comprehensive NLP processing of a document
        
        Args:
            document: Preprocessed document
            
        Returns:
            NLPResults with all analysis completed
        """
        start_time = time.time()
        
        # Generate embeddings
        embeddings = self._generate_embeddings(document.cleaned_text)
        
        # Document classification
        classification = self._classify_document_content(document)
        
        # Named entity recognition
        entities = self._extract_entities(document.cleaned_text)
        
        # Sentiment analysis
        sentiment = self._analyze_sentiment(document.cleaned_text, document.sentences)
        
        # Topic extraction
        topics = self._extract_topics([document.cleaned_text])
        
        # Keyword extraction
        keywords = self._extract_keywords(document.cleaned_text)
        
        # Quality metrics
        quality_metrics = self._assess_quality(document)
        
        processing_time = time.time() - start_time
        
        return NLPResults(
            document_id=document.id,
            embeddings=embeddings,
            classification=classification,
            entities=entities,
            sentiment=sentiment,
            topics=topics,
            keywords=keywords,
            language=document.language,
            quality_metrics=quality_metrics,
            processing_metadata={
                "processing_time": processing_time,
                "model_versions": self._get_model_versions(),
                "timestamp": datetime.now().isoformat()
            }
        )
    
    def _generate_embeddings(self, text: str) -> np.ndarray:
        """Generate semantic embeddings for text"""
        if not text.strip():
            return np.zeros(384)  # Default dimension for MiniLM
        
        try:
            embeddings = self.sentence_transformer.encode(text)
            return embeddings
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            return np.zeros(384)
    
    def _classify_document_content(self, document: ProcessedDocument) -> Dict[str, Any]:
        """Classify document content using multiple approaches"""
        classification = {
            "primary_category": "unknown",
            "confidence": 0.0,
            "subcategories": [],
            "reasoning": {}
        }
        
        try:
            # Use rule-based classification first
            rule_based = self._rule_based_classification(document)
            classification.update(rule_based)
            
            # Add ML-based classification if available
            if hasattr(self, 'classification_model'):
                ml_based = self._ml_based_classification(document)
                classification["ml_prediction"] = ml_based
            
        except Exception as e:
            logger.error(f"Document classification failed: {e}")
        
        return classification
    
    def _rule_based_classification(self, document: ProcessedDocument) -> Dict[str, Any]:
        """Rule-based document classification"""
        text = document.cleaned_text.lower()
        structure = document.structure
        
        # Classification rules
        qa_indicators = text.count('?') / len(text.split()) if text.split() else 0
        
        # Q&A patterns
        qa_patterns = [
            r'\bq\d*[\:\.]', r'\ba\d*[\:\.]', r'question\s*\d*:',
            r'answer\s*\d*:', r'\bfaq\b', r'q&a'
        ]
        qa_score = sum(len(re.findall(pattern, text)) for pattern in qa_patterns)
        
        # Report patterns
        report_patterns = [
            r'\bexecutive\s+summary\b', r'\bintroduction\b', r'\bconclusion\b',
            r'\brecommendations?\b', r'\bmethodology\b', r'\bresults?\b',
            r'\bfindings?\b', r'\banalysis\b'
        ]
        report_score = sum(len(re.findall(pattern, text)) for pattern in report_patterns)
        
        # Essay patterns
        essay_patterns = [
            r'\bin\s+conclusion\b', r'\bto\s+summarize\b', r'\bfirstly\b',
            r'\bsecondly\b', r'\bmoreover\b', r'\bfurthermore\b',
            r'\bhowever\b', r'\bnevertheless\b'
        ]
        essay_score = sum(len(re.findall(pattern, text)) for pattern in essay_patterns)
        
        # Scoring and decision
        scores = {
            "question_answer": qa_score + (qa_indicators * 10),
            "report": report_score + (structure.get("has_headers", 0) * 2),
            "essay": essay_score + (structure.get("paragraph_count", 0) * 0.1)
        }
        
        primary_category = max(scores, key=scores.get)
        max_score = scores[primary_category]
        total_score = sum(scores.values())
        confidence = max_score / total_score if total_score > 0 else 0
        
        return {
            "primary_category": primary_category,
            "confidence": confidence,
            "scores": scores,
            "reasoning": {
                "qa_indicators": qa_indicators,
                "structure_features": structure,
                "pattern_matches": {
                    "qa_patterns": qa_score,
                    "report_patterns": report_score,
                    "essay_patterns": essay_score
                }
            }
        }
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities using spaCy"""
        if not text.strip():
            return []
        
        try:
            doc = self.nlp(text)
            entities = []
            
            for ent in doc.ents:
                entities.append({
                    "text": ent.text,
                    "label": ent.label_,
                    "description": spacy.explain(ent.label_),
                    "start": ent.start_char,
                    "end": ent.end_char,
                    "confidence": getattr(ent, 'confidence', 1.0)
                })
            
            # Remove duplicates and sort by confidence
            unique_entities = {(e["text"], e["label"]): e for e in entities}
            return sorted(unique_entities.values(), 
                         key=lambda x: x["confidence"], reverse=True)
            
        except Exception as e:
            logger.error(f"Entity extraction failed: {e}")
            return []
    
    def _analyze_sentiment(self, text: str, sentences: List[str]) -> Dict[str, Any]:
        """Comprehensive sentiment analysis"""
        if not text.strip():
            return {"overall": "neutral", "confidence": 0.0}
        
        try:
            sentiment_result = {
                "overall": "neutral",
                "confidence": 0.0,
                "scores": {},
                "sentence_sentiments": []
            }
            
            # Overall sentiment using VADER
            vader_scores = self.sentiment_analyzer.polarity_scores(text)
            sentiment_result["scores"]["vader"] = vader_scores
            
            # Transformer-based sentiment
            try:
                transformer_result = self.sentiment_model(text[:512])  # Limit length
                sentiment_result["scores"]["transformer"] = transformer_result[0]
            except Exception as e:
                logger.warning(f"Transformer sentiment failed: {e}")
            
            # Sentence-level sentiment
            for sentence in sentences[:10]:  # Limit to first 10 sentences
                sent_sentiment = self.sentiment_analyzer.polarity_scores(sentence)
                sentiment_result["sentence_sentiments"].append({
                    "sentence": sentence[:100],  # Truncate for storage
                    "sentiment": sent_sentiment
                })
            
            # Determine overall sentiment
            compound_score = vader_scores['compound']
            if compound_score >= 0.05:
                sentiment_result["overall"] = "positive"
                sentiment_result["confidence"] = compound_score
            elif compound_score <= -0.05:
                sentiment_result["overall"] = "negative"
                sentiment_result["confidence"] = abs(compound_score)
            else:
                sentiment_result["overall"] = "neutral"
                sentiment_result["confidence"] = 1 - abs(compound_score)
            
            return sentiment_result
            
        except Exception as e:
            logger.error(f"Sentiment analysis failed: {e}")
            return {"overall": "neutral", "confidence": 0.0, "error": str(e)}
    
    def _extract_topics(self, texts: List[str]) -> List[Dict[str, Any]]:
        """Extract topics using BERTopic"""
        if not texts or not any(text.strip() for text in texts):
            return []
        
        try:
            # Filter out empty texts
            valid_texts = [text for text in texts if text.strip()]
            
            if len(valid_texts) < 2:
                # Use keyword extraction for single documents
                return self._extract_topics_keywords(valid_texts[0])
            
            topics, probabilities = self.topic_model.fit_transform(valid_texts)
            
            topic_results = []
            for topic_id in set(topics):
                if topic_id != -1:  # Skip outlier topic
                    topic_info = self.topic_model.get_topic(topic_id)
                    topic_results.append({
                        "topic_id": topic_id,
                        "keywords": [word for word, score in topic_info[:10]],
                        "keyword_scores": dict(topic_info[:10]),
                        "document_count": list(topics).count(topic_id)
                    })
            
            return sorted(topic_results, key=lambda x: x["document_count"], reverse=True)
            
        except Exception as e:
            logger.error(f"Topic extraction failed: {e}")
            return self._extract_topics_keywords(texts[0] if texts else "")
    
    def _extract_topics_keywords(self, text: str) -> List[Dict[str, Any]]:
        """Fallback topic extraction using keywords"""
        if not text.strip():
            return []
        
        try:
            keywords = self.keyword_extractor.extract_keywords(text)
            return [{
                "topic_id": 0,
                "keywords": [kw[1] for kw in keywords[:10]],
                "keyword_scores": {kw[1]: kw[0] for kw in keywords[:10]},
                "document_count": 1
            }]
        except Exception as e:
            logger.error(f"Keyword topic extraction failed: {e}")
            return []
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract keywords using YAKE"""
        if not text.strip():
            return []
        
        try:
            keywords = self.keyword_extractor.extract_keywords(text)
            return [kw[1] for kw in keywords]
        except Exception as e:
            logger.error(f"Keyword extraction failed: {e}")
            return []
    
    def _assess_quality(self, document: ProcessedDocument) -> Dict[str, Any]:
        """Assess document quality using multiple metrics"""
        quality_metrics = {
            "readability": {},
            "coherence": 0.0,
            "completeness": 0.0,
            "overall_score": 0.0
        }
        
        try:
            text = document.cleaned_text
            
            # Readability metrics
            if text.strip():
                quality_metrics["readability"] = {
                    "flesch_reading_ease": flesch_reading_ease(text),
                    "flesch_kincaid_grade": flesch_kincaid_grade(text),
                    "automated_readability_index": automated_readability_index(text)
                }
            
            # Coherence (simple version using sentence similarity)
            sentences = document.sentences
            if len(sentences) > 1:
                embeddings = [self.sentence_transformer.encode(sent) for sent in sentences[:10]]
                similarities = []
                for i in range(len(embeddings) - 1):
                    sim = cosine_similarity([embeddings[i]], [embeddings[i + 1]])[0][0]
                    similarities.append(sim)
                quality_metrics["coherence"] = np.mean(similarities) if similarities else 0.0
            
            # Completeness (based on structure)
            structure = document.structure
            completeness_score = 0.0
            if structure.get("sentence_count", 0) > 3:
                completeness_score += 0.3
            if structure.get("paragraph_count", 0) > 1:
                completeness_score += 0.3
            if structure.get("word_count", 0) > 100:
                completeness_score += 0.4
                
            quality_metrics["completeness"] = completeness_score
            
            # Overall score (weighted average)
            readability_score = (quality_metrics["readability"].get("flesch_reading_ease", 0) / 100)
            quality_metrics["overall_score"] = (
                readability_score * 0.4 + 
                quality_metrics["coherence"] * 0.3 + 
                quality_metrics["completeness"] * 0.3
            )
            
        except Exception as e:
            logger.error(f"Quality assessment failed: {e}")
        
        return quality_metrics
    
    def _get_model_versions(self) -> Dict[str, str]:
        """Get versions of loaded models"""
        return {
            "sentence_transformer": "all-MiniLM-L6-v2",
            "spacy": spacy.__version__,
            "transformers": "latest",
            "bertopic": "latest"
        }


class DocumentClassifier:
    """Specialized document classifier for essay/report/Q&A classification"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self